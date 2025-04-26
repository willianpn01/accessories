from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import os
from pdfplumber import open as open_pdf

def docx_para_pdf(docx_path, pdf_path):
    doc = Document(docx_path)
    c = canvas.Canvas(pdf_path, pagesize=A4)
    largura, altura = A4
    y = altura - 40
    for para in doc.paragraphs:
        texto = para.text
        if texto.strip() == '':
            y -= 18
            continue
        c.drawString(40, y, texto)
        y -= 18
        if y < 40:
            c.showPage()
            y = altura - 40
    c.save()
    return pdf_path

def pdf_para_docx(pdf_path, docx_path):
    import docx
    doc = docx.Document()
    with open_pdf(pdf_path) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ''
            for linha in texto.splitlines():
                doc.add_paragraph(linha)
            doc.add_page_break()
    doc.save(docx_path)
    return docx_path
